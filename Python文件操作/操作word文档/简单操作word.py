# -*- coding: utf-8 -*-
"""
pip install python-docx
使用：
import docx
可能是原来就有docx 这个 模块，所以这个起名为 python-docx
"""

import docx

# opening a new document
document = docx.Document()

paragraph = document.add_paragraph('作者：苏轼', 'Subtitle')
prior_paragraph = paragraph.insert_paragraph_before('水调歌头', 'Title')
document.add_paragraph('''
明月几时有？把酒问青天。
不知天上宫阙，今夕是何年。
我欲乘风归去，又恐琼楼玉宇，高处不胜寒。
起舞弄清影，何似在人间。
转朱阁，低绮户，照无眠。
不应有恨，何事长向别时圆？
人有悲欢离合，月有阴晴圆缺，此事古难全。
但愿人长久，千里共婵娟。
''', 'Body Text 2')

# adding a heading
document.add_heading('The real meaning of the universe')
document.add_heading('Heading 0', level=0)
document.add_heading('Heading 1', level=1)
document.add_heading('Heading 2', level=2)

for p in document.paragraphs:
    print('len(p.text):', len(p.text))
    print('p.style.name:', p.style.name)


from docx.enum.style import WD_STYLE_TYPE

document.save('test.docx')

